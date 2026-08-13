"""Semantic and OOXML audits for the Beast Road GDD v4 deliverables."""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.oxml.ns import qn


def fail(message: str) -> None:
    print(f"FAIL: {message}")
    raise SystemExit(1)


def audit_markdown(path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    required = [
        "# Part I - Product Vision",
        "# Part XIV - Reconciliation and Changelog",
        "## 15. Command: The Anti-Idle Combat Loop",
        "## 36. Sigils and Legacy Rank",
        "## 51. Production Milestones",
        "## 52. Release Acceptance Checklist",
        "## 54. Explicitly Out of Scope for 1.0",
        "## 57. v4 Locked Decisions Summary",
    ]
    for phrase in required:
        if phrase not in text:
            fail(f"missing required section: {phrase}")

    if len(re.findall(r"^# Part ", text, flags=re.MULTILINE)) != 14:
        fail("expected 14 major parts")
    if len(re.findall(r"^## \d+\.", text, flags=re.MULTILINE)) != 58:
        fail("expected numbered sections 0 through 57")
    checklist_count = len(re.findall(r"^- \[ \]", text, flags=re.MULTILINE))
    if checklist_count < 50:
        fail("production/release checklist is unexpectedly short")

    contradictory = [
        ("Air is the only starting element", "obsolete element gating"),
        ("ten victories", "obsolete ten-clear grind"),
        ("fifteen crossroads", "obsolete crossroad count"),
        ("Tusken Raiders", "copyrighted shorthand"),
    ]
    for phrase, meaning in contradictory:
        if phrase.lower() in text.lower():
            fail(f"found {meaning}: {phrase}")

    if "three acts plus one final summit".lower() not in text.lower():
        fail("run structure summary is missing")
    if "all four elements are available".lower() not in text.lower():
        fail("early element availability is missing")
    if "third lethal down".lower() not in text.lower():
        fail("hero defeat rule is missing")

    print(
        "Markdown audit: PASS | "
        f"chars={len(text)} parts=14 numbered_sections=58 "
        f"checklist_items={checklist_count}"
    )


def audit_docx(path: Path) -> None:
    doc = Document(path)
    body_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        body_text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)

    if "BEAST ROAD" not in body_text or "Final Creative Standard" not in body_text:
        fail("DOCX lost beginning or ending content")
    if len(doc.tables) < 14:
        fail("DOCX table count is unexpectedly low")
    heading_counts = {
        style: sum(1 for p in doc.paragraphs if p.style.name == style)
        for style in ("Heading 1", "Heading 2", "Heading 3")
    }
    if heading_counts != {"Heading 1": 15, "Heading 2": 59, "Heading 3": 88}:
        fail(f"unexpected heading structure: {heading_counts}")

    for table_index, table in enumerate(doc.tables, start=1):
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.first_child_found_in("w:tblW")
        layout = tbl_pr.first_child_found_in("w:tblLayout")
        if tbl_w is None or tbl_w.get(qn("w:w")) != "9360":
            fail(f"table {table_index} does not use 9360 DXA width")
        if layout is None or layout.get(qn("w:type")) != "fixed":
            fail(f"table {table_index} does not use fixed layout")
        grid = table._tbl.tblGrid
        grid_widths = [int(col.get(qn("w:w"))) for col in grid]
        if sum(grid_widths) != 9360:
            fail(f"table {table_index} grid does not sum to 9360 DXA")
        for row in table.rows:
            if len(row.cells) != len(grid_widths):
                fail(f"table {table_index} contains merged/irregular cells")
            for idx, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if tc_w is None or int(tc_w.get(qn("w:w"))) != grid_widths[idx]:
                    fail(f"table {table_index} cell width mismatch")
        tr_pr = table.rows[0]._tr.get_or_add_trPr()
        if tr_pr.find(qn("w:tblHeader")) is None:
            fail(f"table {table_index} lacks a repeated header row")

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(path) as archive:
        styles = ET.fromstring(archive.read("word/styles.xml"))
        document = ET.fromstring(archive.read("word/document.xml"))
        numbering = ET.fromstring(archive.read("word/numbering.xml"))
    title_style = styles.find(".//w:style[@w:styleId='Title']", ns)
    if title_style is not None and title_style.find(".//w:pBdr", ns) is not None:
        fail("Title style retains border residue")
    first_paragraphs = document.findall(".//w:body/w:p", ns)[:3]
    for paragraph in first_paragraphs:
        if paragraph.find(".//w:pBdr", ns) is not None:
            fail("leading title block contains border residue")
    if len(numbering.findall(".//w:abstractNum", ns)) < 3:
        fail("real bullet, number, and checklist numbering definitions are missing")

    print(
        "DOCX audit: PASS | "
        f"paragraphs={len(doc.paragraphs)} tables={len(doc.tables)} "
        f"headings={heading_counts} chars={len(body_text)}"
    )


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: audit_gdd_v4.py SOURCE.md OUTPUT.docx")
    audit_markdown(Path(sys.argv[1]))
    audit_docx(Path(sys.argv[2]))
