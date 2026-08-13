"""Build the Google Docs-ready Beast Road GDD v4 DOCX from its Markdown source."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH,
    WD_BREAK,
    WD_LINE_SPACING,
    WD_TAB_ALIGNMENT,
)
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


FONT = "Arial"
INK = "20292C"
NAVY = "21363E"
EMBER = "B85D2C"
SLATE = "53656A"
MUTED = "68787C"
SAND = "F3ECE4"
ICE = "EDF3F3"
PAPER = "FBFCFB"
BORDER = "C8D2D3"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
BODY_SIZE = 10.5


def set_run_font(run, *, size: float = BODY_SIZE, bold: bool | None = None,
                 italic: bool | None = None, color: str = INK,
                 font: str = FONT) -> None:
    run.font.name = font
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_spacing(paragraph, *, before: float = 0, after: float = 8,
                line: float = 1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line
    fmt.widow_control = True


def set_cell_margins(cell, top: int = 80, start: int = 120,
                     bottom: int = 80, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start),
                        ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def keep_table_row_together(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, *, edge: str, color: str,
                         size: int = 8, space: int = 8) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    border = p_bdr.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        p_bdr.append(border)
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), str(space))
    border.set(qn("w:color"), color)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths: list[int], *, indent: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    layout = tbl_pr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def choose_widths(rows: list[list[str]]) -> list[int]:
    cols = max(len(row) for row in rows)
    headers = [cell.strip().lower() for cell in rows[0]]
    if headers == ["resource", "persistence", "sources", "sinks"]:
        return [1350, 1560, 3225, 3225]

    max_lengths = []
    for col in range(cols):
        values = [len(re.sub(r"[*`]", "", row[col])) if col < len(row) else 0
                  for row in rows]
        max_lengths.append(max(8, min(max(values), 80)))

    if cols == 2:
        first = max_lengths[0]
        second = max_lengths[1]
        ratio = first / max(first + second, 1)
        first_width = int(TABLE_WIDTH_DXA * min(0.40, max(0.25, ratio)))
        return [first_width, TABLE_WIDTH_DXA - first_width]

    minimum = 1260 if cols == 3 else 1120 if cols == 4 else 780
    remaining = TABLE_WIDTH_DXA - minimum * cols
    weights = [max(3.0, length ** 0.5) for length in max_lengths]
    total = sum(weights)
    widths = [minimum + int(remaining * weight / total) for weight in weights]

    scale = TABLE_WIDTH_DXA / sum(widths)
    widths = [int(w * scale) for w in widths]
    widths[-1] += TABLE_WIDTH_DXA - sum(widths)
    return widths


INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|(?<!\*)\*[^*]+\*(?!\*))")


def add_inline(paragraph, text: str, *, size: float = BODY_SIZE,
               base_bold: bool = False, color: str = INK) -> None:
    position = 0
    for match in INLINE.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            if base_bold or size != BODY_SIZE or color != INK:
                set_run_font(run, size=size, bold=base_bold, color=color)
        token = match.group(0)
        if token.startswith("**"):
            label = token[2:-2]
            decision_color = {
                "LOCKED": NAVY,
                "TUNE": EMBER,
                "OPEN": "9B4B2A",
                "POST-LAUNCH": SLATE,
                "CUT": "7A4A4A",
            }.get(label.rstrip(":").upper(), color)
            run = paragraph.add_run(label)
            set_run_font(run, size=size, bold=True, color=decision_color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=max(9.5, size - 0.5), bold=base_bold,
                         color=MUTED, font="Consolas")
        else:
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size, bold=base_bold, italic=True, color=color)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        if base_bold or size != BODY_SIZE or color != INK:
            set_run_font(run, size=size, bold=base_bold, color=color)


def add_numbering_abstract(document, *, kind: str) -> int:
    numbering = document.part.numbering_part.element
    existing = [int(x.get(qn("w:abstractNumId")))
                for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing, default=-1) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "decimal" if kind == "number" else "bullet")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if kind == "number":
            marker = f"%{level + 1}."
        elif kind == "check":
            marker = "☐"
        else:
            marker = ("●", "○", "■")[level]
        lvl_text.set(qn("w:val"), marker)
        lvl.append(lvl_text)
        jc = OxmlElement("w:lvlJc")
        jc.set(qn("w:val"), "left")
        lvl.append(jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(720 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(720 + level * 360))
        ind.set(qn("w:hanging"), "360")
        p_pr.append(ind)
        lvl.append(p_pr)
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), FONT)
        fonts.set(qn("w:hAnsi"), FONT)
        r_pr.append(fonts)
        lvl.append(r_pr)
        abstract.append(lvl)
    numbering.append(abstract)
    return abstract_id


def add_num_instance(document, abstract_id: int) -> int:
    numbering = document.part.numbering_part.element
    existing = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    # LibreOffice otherwise continues decimal lists that share an abstract
    # definition, even when each Markdown list receives a fresh numId.
    level_override = OxmlElement("w:lvlOverride")
    level_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    level_override.append(start_override)
    num.append(level_override)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int, level: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(level, 2)))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


def configure_styles(document) -> None:
    styles = document.styles
    title = styles["Title"]
    if title._element.pPr is not None:
        title_border = title._element.pPr.find(qn("w:pBdr"))
        if title_border is not None:
            title._element.pPr.remove(title_border)

    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.widow_control = True

    tokens = {
        "Heading 1": (18.5, NAVY, 18, 7),
        "Heading 2": (14.5, NAVY, 16, 5),
        "Heading 3": (11.5, SLATE, 12, 4),
    }
    for name, (size, color, before, after) in tokens.items():
        style = styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "GDD Quote" not in styles:
        quote = styles.add_style("GDD Quote", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["GDD Quote"]
    quote.font.name = FONT
    quote._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    quote._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    quote.font.size = Pt(11)
    quote.font.italic = True
    quote.font.color.rgb = RGBColor.from_string(NAVY)
    quote.paragraph_format.left_indent = Inches(0.25)
    quote.paragraph_format.right_indent = Inches(0.25)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(12)
    quote.paragraph_format.line_spacing = 1.15

    if "GDD Exit Gate" not in styles:
        gate = styles.add_style("GDD Exit Gate", WD_STYLE_TYPE.PARAGRAPH)
    else:
        gate = styles["GDD Exit Gate"]
    gate.font.name = FONT
    gate._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    gate._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    gate.font.size = Pt(10)
    gate.font.color.rgb = RGBColor.from_string(NAVY)
    gate.paragraph_format.left_indent = Inches(0.16)
    gate.paragraph_format.right_indent = Inches(0.12)
    gate.paragraph_format.space_before = Pt(6)
    gate.paragraph_format.space_after = Pt(9)
    gate.paragraph_format.line_spacing = 1.1
    gate.paragraph_format.keep_together = True


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c) for c in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(document, rows: list[list[str]]) -> None:
    cols = max(len(row) for row in rows)
    normalized = [row + [""] * (cols - len(row)) for row in rows]
    table = document.add_table(rows=len(normalized), cols=cols)
    set_table_geometry(table, choose_widths(normalized))
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    centered_headers = {
        "act", "cap", "count", "cost", "duration", "mode", "owner",
        "phase", "priority", "road", "slot", "status", "tier", "uses",
        "value", "wave",
    }
    for row_idx, row in enumerate(normalized):
        keep_table_row_together(table.rows[row_idx])
        for col_idx, value in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            shade_cell(cell, NAVY if row_idx == 0 else (PAPER if row_idx % 2 else ICE))
            paragraph = cell.paragraphs[0]
            set_spacing(paragraph, before=0, after=0, line=1.08)
            font_size = 8.5 if cols >= 5 else 9
            add_inline(
                paragraph,
                value,
                size=font_size,
                base_bold=row_idx == 0,
                color=WHITE if row_idx == 0 else INK,
            )
            header = normalized[0][col_idx].strip().lower()
            if header in centered_headers:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after = document.add_paragraph()
    set_spacing(after, after=4)


def add_cover_snapshot(document) -> None:
    entries = (
        ("ONE HERO", "Field commander"),
        ("FOUR ROADS", "Concurrent threats"),
        ("THREE ACTS", "Plus final summit"),
        ("55–65 MIN", "Target run"),
    )
    table = document.add_table(rows=1, cols=4)
    set_table_geometry(table, [TABLE_WIDTH_DXA // 4] * 4, indent=80)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_repeat_table_header(table.rows[0])
    keep_table_row_together(table.rows[0])
    for idx, (label, detail) in enumerate(entries):
        cell = table.cell(0, idx)
        shade_cell(cell, NAVY)
        set_cell_margins(cell, top=100, start=80, bottom=100, end=80)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, after=1, line=1)
        run = p.add_run(label)
        set_run_font(run, size=9, bold=True, color=WHITE)
        p.add_run().add_break()
        run = p.add_run(detail)
        set_run_font(run, size=7.5, color="DCE6E7")
    after = document.add_paragraph()
    set_spacing(after, after=3)


def add_cover(document, lines: list[str], art_path: Path) -> int:
    title = lines[0].removeprefix("# ").strip()
    subtitle = lines[2].removeprefix("## ").strip()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=5)
    run = p.add_run("OFFICIAL PRODUCTION MASTER")
    set_run_font(run, size=8.5, bold=True, color=EMBER)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=1)
    run = p.add_run(title)
    set_run_font(run, size=31, bold=True, color=NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=10)
    run = p.add_run(subtitle)
    set_run_font(run, size=14, color=SLATE)

    if art_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=0, after=7)
        shape = p.add_run().add_picture(str(art_path), width=Inches(6.5))
        shape._inline.docPr.set("title", "Beast Road key art")
        shape._inline.docPr.set(
            "descr",
            "A moonlit colossal walking beast carries a warm firelit settlement through "
            "the wilderness, establishing the scale and tone of Beast Road.",
        )

    add_cover_snapshot(document)

    idx = 4
    metadata: list[str] = []
    while idx < len(lines) and lines[idx].strip() != "---":
        raw = lines[idx].strip()
        if raw and not raw.startswith(">"):
            metadata.append(raw.rstrip("  "))
        idx += 1

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=2, after=2)
    add_inline(p,
        "One hero, four roads, and a moving town that cannot be defended everywhere at once. "
        "Towers hold the formation. The hero, the route, and the player's timing decide where "
        "the defense bends and where it breaks.", size=10.5, base_bold=True, color=NAVY)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=5, after=0, line=1.05)
    compact = " • ".join(
        re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
        for item in metadata[:2]
    )
    add_inline(p, compact, size=8.5, color=MUTED)
    if len(metadata) > 2:
        p.add_run().add_break()
        add_inline(p, " • ".join(
            re.sub(r"\*\*([^*]+)\*\*", r"\1", item)
            for item in metadata[2:]
        ), size=7.5, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)
    return idx + 1


def add_contents(document, lines: list[str]) -> None:
    h = document.add_paragraph("Contents", style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    shade_paragraph(h, SAND)
    set_paragraph_border(h, edge="left", color=EMBER, size=14, space=8)
    parts = [line[2:].strip() for line in lines if line.startswith("# Part ")]
    for part in parts:
        p = document.add_paragraph()
        set_spacing(p, after=4, line=1.05)
        match = re.match(r"Part ([IVX]+)\s+-\s+(.+)", part)
        if match:
            run = p.add_run(f"PART {match.group(1)}")
            set_run_font(run, size=8.5, bold=True, color=EMBER)
            run = p.add_run(f"   {match.group(2)}")
            set_run_font(run, size=10, bold=True, color=NAVY)
        else:
            add_inline(p, part, size=10, base_bold=True, color=NAVY)
    p = document.add_paragraph()
    set_spacing(p, before=10, after=0)
    add_inline(p,
        "Google Docs and Word will expose every numbered section in the document outline.",
        size=9, color=MUTED)
    p.add_run().add_break(WD_BREAK.PAGE)


def add_closing_plate(document, art_path: Path) -> None:
    if art_path.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_spacing(p, before=14, after=6)
        shape = p.add_run().add_picture(str(art_path), width=Inches(6.5))
        shape._inline.docPr.set("title", "Crown of the World concept art")
        shape._inline.docPr.set(
            "descr",
            "A monumental dark summit rises through ember-lit fog, representing the final "
            "ascent and the production destination for Beast Road.",
        )
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=2, after=0)
    run = p.add_run("CLARITY  •  PRESSURE  •  PERSONALITY  •  CONSEQUENCE  •  PAYOFF")
    set_run_font(run, size=8.5, bold=True, color=EMBER)


def build(source: Path, output: Path) -> None:
    lines = source.read_text(encoding="utf-8").splitlines()
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True
    configure_styles(document)

    header = section.header.paragraphs[0]
    header.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    set_spacing(header, after=0, line=1)
    run = header.add_run("BEAST ROAD  /  GDD v4.0")
    set_run_font(run, size=7.5, bold=True, color=SLATE)
    header.add_run("\t")
    run = header.add_run("AUTHORITATIVE 1.0 DESIGN")
    set_run_font(run, size=7.5, bold=True, color=EMBER)

    footer = section.footer.paragraphs[0]
    footer.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), alignment=WD_TAB_ALIGNMENT.RIGHT
    )
    set_spacing(footer, after=0, line=1)
    run = footer.add_run("INTERNAL PRODUCTION MASTER")
    set_run_font(run, size=7.5, color=MUTED)
    footer.add_run("\t")
    run = footer.add_run("PAGE ")
    set_run_font(run, size=7.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    field_run = OxmlElement("w:r")
    field_rpr = OxmlElement("w:rPr")
    field_color = OxmlElement("w:color")
    field_color.set(qn("w:val"), MUTED)
    field_size = OxmlElement("w:sz")
    field_size.set(qn("w:val"), "15")
    field_rpr.append(field_color)
    field_rpr.append(field_size)
    field_run.append(field_rpr)
    field_text = OxmlElement("w:t")
    field_text.text = "1"
    field_run.append(field_text)
    field.append(field_run)
    footer._p.append(field)

    core = document.core_properties
    core.title = "Beast Road - Game Design Document v4.0"
    core.subject = "Authoritative production target for Beast Road 1.0"
    core.author = "Beast Road Team"
    core.keywords = "game design, tower defense, roguelite, production"
    core.comments = "Generated from docs/Game_Design_v4.md"

    index = add_cover(
        document,
        lines,
        source.parent.parent / "game" / "art" / "bg" / "menu_key_art.png",
    )
    add_contents(document, lines)

    bullet_abs = add_numbering_abstract(document, kind="bullet")
    number_abs = add_numbering_abstract(document, kind="number")
    check_abs = add_numbering_abstract(document, kind="check")
    active_list: tuple[str, int] | None = None

    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()

        if not stripped or stripped == "---":
            active_list = None
            index += 1
            continue

        if stripped.startswith("|"):
            rows, index = parse_table(lines, index)
            add_table(document, rows)
            active_list = None
            continue

        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            text = heading.group(2)
            if level == 1 and text.startswith("Part "):
                pass
            p = document.add_paragraph(style=f"Heading {level}")
            add_inline(p, text)
            if level == 1:
                shade_paragraph(p, SAND)
                set_paragraph_border(p, edge="left", color=EMBER, size=14, space=8)
            active_list = None
            index += 1
            continue

        if stripped.startswith(">"):
            text = stripped.lstrip("> ")
            p = document.add_paragraph(style="GDD Quote")
            shade_paragraph(p, SAND)
            set_paragraph_border(p, edge="left", color=EMBER, size=9, space=7)
            add_inline(p, text, size=11, color=NAVY)
            active_list = None
            index += 1
            continue

        checkbox = re.match(r"^(\s*)- \[ \]\s+(.+)$", raw)
        bullet = re.match(r"^(\s*)-\s+(.+)$", raw)
        numbered = re.match(r"^(\s*)\d+\.\s+(.+)$", raw)
        match = checkbox or bullet or numbered
        if match:
            if checkbox:
                kind, abstract = "check", check_abs
            elif numbered:
                kind, abstract = "number", number_abs
            else:
                kind, abstract = "bullet", bullet_abs
            level = min(len(match.group(1).expandtabs(4)) // 2, 2)
            if active_list is None or active_list[0] != kind:
                active_list = (kind, add_num_instance(document, abstract))
            p = document.add_paragraph()
            set_spacing(p, after=4)
            apply_numbering(p, active_list[1], level)
            add_inline(p, match.group(2), size=BODY_SIZE)
            index += 1
            continue

        is_exit_gate = stripped.startswith("**Exit gate:**")
        p = document.add_paragraph(style="GDD Exit Gate" if is_exit_gate else None)
        if is_exit_gate:
            shade_paragraph(p, ICE)
            set_paragraph_border(p, edge="left", color=NAVY, size=10, space=7)
        else:
            set_spacing(p, after=6)
        add_inline(p, stripped.rstrip("  "), size=10 if is_exit_gate else BODY_SIZE)
        active_list = None
        index += 1

    add_closing_plate(
        document,
        source.parent.parent / "game" / "art" / "bg" / "macro_act3.png",
    )
    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: build_gdd_v4_docx.py SOURCE.md OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
